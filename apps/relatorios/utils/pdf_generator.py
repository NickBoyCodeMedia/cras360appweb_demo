"""
Módulo para gerar arquivos PDF para relatórios e documentos instrumentais.
"""

import os
import tempfile
from datetime import datetime
from io import BytesIO
import base64
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm, inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.pdfgen import canvas
from PIL import Image as PILImage
from apps.core.models import CRAS
# Importe Cadastro e Usuario se existirem, ou ajuste conforme necessário:
# from apps.core.models import Cadastro
# from apps.auth_app.models import Usuario

# Ajuste a importação do config conforme o layout do seu projeto.
# Exemplo para importação absoluta (ajuste se necessário):
try:
    from config import RELATORIOS_DIR, PDF_GERADOS_DIR
except ImportError:
    # Tente importar usando caminho relativo ao projeto Django/Flask
    try:
        from apps.relatorios.utils import config
        RELATORIOS_DIR = config.RELATORIOS_DIR
        PDF_GERADOS_DIR = config.PDF_GERADOS_DIR
    except ImportError:
        # Se ainda falhar, defina variáveis dummy ou lance erro
        RELATORIOS_DIR = None
        PDF_GERADOS_DIR = None
        # raise ImportError("Não foi possível importar config. Ajuste o caminho de importação.")

# Configurações padrão para documentos
PAGE_WIDTH, PAGE_HEIGHT = A4
MARGIN_LEFT = 2.5 * cm
MARGIN_RIGHT = 2.0 * cm
MARGIN_TOP = 3.0 * cm
MARGIN_BOTTOM = 2.5 * cm

def gerar_pdf_relatorio_formularios(cadastros, colunas, titulo=None):
    """
    Gera um relatório PDF a partir dos cadastros filtrados.
    """
    # Criar um arquivo temporário para o PDF
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    arquivo_pdf = temp_file.name
    temp_file.close()
    
    # Configurar o documento
    doc = SimpleDocTemplate(
        arquivo_pdf,
        pagesize=A4,
        leftMargin=MARGIN_LEFT,
        rightMargin=MARGIN_RIGHT,
        topMargin=MARGIN_TOP,
        bottomMargin=MARGIN_BOTTOM
    )
    
    # Preparar estilos
    styles = getSampleStyleSheet()
    titulo_style = styles['Heading1']
    normal_style = styles['Normal']
    
    # Elementos para o documento
    elementos = []
    
    # Adicionar cabeçalho
    elementos.append(Paragraph("Cabeçalho Personalizado - Relatório", titulo_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar título
    titulo_texto = titulo or f"Relatório de Cadastros - {datetime.now().strftime('%d/%m/%Y')}"
    elementos.append(Paragraph(titulo_texto, titulo_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar informações do relatório
    elementos.append(Paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}", normal_style))
    elementos.append(Paragraph(f"Total de registros: {len(cadastros)}", normal_style))
    elementos.append(Spacer(1, 1 * cm))
    
    # Preparar cabeçalho da tabela
    cabecalho = []
    colunas_unicas = list(set(colunas))  # Garantir que as colunas não se repitam
    for coluna in colunas_unicas:
        if coluna == 'nome':
            cabecalho.append('Nome')
        elif coluna == 'projeto':
            cabecalho.append('Projeto')
        elif coluna == 'data_nascimento':
            cabecalho.append('Data Nasc.')
        elif coluna == 'data_cadastro':
            cabecalho.append('Data Cadastro')
        elif coluna == 'cras':
            cabecalho.append('CRAS')
        elif coluna == 'cidade':
            cabecalho.append('Cidade')
        elif coluna == 'bairro':
            cabecalho.append('Bairro')
        elif coluna == 'contato':
            cabecalho.append('Contato')
        elif coluna == 'nome_mae':
            cabecalho.append('Nome da Mãe')
    
    # Preparar dados da tabela
    dados_tabela = [cabecalho]
    for cadastro in cadastros:
        linha = []
        for coluna in colunas_unicas:  # Usar colunas únicas
            if coluna == 'nome':
                linha.append(cadastro.nome or '')
            elif coluna == 'projeto':
                linha.append(cadastro.projeto or '')
            elif coluna == 'data_nascimento':
                linha.append(cadastro.data_nascimento.strftime('%d/%m/%Y') if cadastro.data_nascimento else '')
            elif coluna == 'data_cadastro':
                linha.append(cadastro.data_cadastro.strftime('%d/%m/%Y') if cadastro.data_cadastro else '')
            elif coluna == 'cras':
                linha.append(cadastro.cras.nome if cadastro.cras else '')
            elif coluna == 'cidade':
                linha.append(cadastro.cidade or '')
            elif coluna == 'bairro':
                linha.append(cadastro.bairro or '')
            elif coluna == 'contato':
                linha.append(cadastro.contato or '')
            elif coluna == 'nome_mae':
                linha.append(cadastro.nome_mae or '')
        dados_tabela.append(linha)
    
    # Criar tabela
    tabela = Table(dados_tabela, repeatRows=1)
    
    # Estilo da tabela
    estilo_tabela = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ])
    
    # Aplicar estilo alternado às linhas
    for i in range(1, len(dados_tabela)):
        if i % 2 == 0:
            estilo_tabela.add('BACKGROUND', (0, i), (-1, i), colors.lightgrey)
    
    tabela.setStyle(estilo_tabela)
    elementos.append(tabela)
    
    # Adicionar rodapé
    elementos.append(Spacer(1, 1 * cm))
    elementos.append(Paragraph("Relatório gerado pelo Sistema CRAS360 - © 2025", normal_style))
    
    # Gerar o documento
    doc.build(elementos)
    
    # Retornar o arquivo para download
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"  # ou .docx/.xlsx conforme o caso
    return django_send_file(arquivo_pdf, download_name, 'application/pdf')

def gerar_pdf_instrumental(titulo, conteudo=None):
    """
    Gera um documento PDF instrumental com título e conteúdo
    
    Args:
        titulo: Título do documento
        conteudo: Conteúdo do documento (pode ser Delta JSON do Quill)
    
    Returns:
        Response com o arquivo PDF para download
    """
    # Criar um arquivo temporário para o PDF
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    arquivo_pdf = temp_file.name
    temp_file.close()
    
    # Configurar o documento
    doc = SimpleDocTemplate(
        arquivo_pdf,
        pagesize=A4,
        leftMargin=MARGIN_LEFT,
        rightMargin=MARGIN_RIGHT,
        topMargin=MARGIN_TOP,
        bottomMargin=MARGIN_BOTTOM,
        title=titulo
    )
    
    # Preparar estilos
    styles = getSampleStyleSheet()
    titulo_style = styles['Heading1']
    subtitulo_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Elementos para o documento
    elementos = []
    
    # Adicionar título
    elementos.append(Paragraph(titulo, titulo_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Processar conteúdo (pode ser Delta JSON do Quill)
    if conteudo:
        import json
        try:
            # Tentar interpretar como delta JSON
            delta = json.loads(conteudo)
            elementos.extend(processar_delta_quill(delta))
        except:
            # Se falhar, tratar como texto simples
            paragrafos = conteudo.split('\n')
            for paragrafo in paragrafos:
                if paragrafo.strip():  # evitar adicionar parágrafos vazios
                    elementos.append(Paragraph(paragrafo, normal_style))
                    elementos.append(Spacer(1, 0.25 * cm))
    
    # Adicionar rodapé
    elementos.append(Spacer(1, 1 * cm))
    elementos.append(Paragraph(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}", normal_style))
    
    # Gerar o documento
    doc.build(elementos)
    
    # Retornar o arquivo para download
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return django_send_file(arquivo_pdf, download_name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def gerar_pdf_relatorio_com_template(cadastros, colunas, template_id, titulo=None):
    """
    Gera um relatório PDF usando um template personalizado
    
    Args:
        cadastros: Lista de objetos Cadastro
        colunas: Lista de colunas a exibir
        template_id: ID do template a ser usado
        titulo: Título opcional para o relatório
    
    Returns:
        Response com o arquivo PDF para download
    """
    from models import UsuarioTemplate
    
    # Obter o template
    template = UsuarioTemplate.query.get(template_id)
    if not template:
        return gerar_pdf_relatorio_formularios(cadastros, colunas, titulo)
    
    # Criar um arquivo temporário para o PDF
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    arquivo_pdf = temp_file.name
    temp_file.close()
    
    # Configurar o documento com as margens do template
    doc = SimpleDocTemplate(
        arquivo_pdf,
        pagesize=A4,
        leftMargin=float(template.margem_esquerda or MARGIN_LEFT),
        rightMargin=float(template.margem_direita or MARGIN_RIGHT),
        topMargin=float(template.margem_superior or MARGIN_TOP),
        bottomMargin=float(template.margem_inferior or MARGIN_BOTTOM)
    )
    
    # Preparar estilos personalizados baseados no template
    styles = getSampleStyleSheet()
    
    titulo_style = ParagraphStyle(
        'TituloPersonalizado',
        parent=styles['Heading1'],
        fontName=template.fonte_titulo or 'Helvetica-Bold',
        fontSize=int(template.tamanho_titulo or 18),
        alignment=1 if template.cabecalho_centralizado else 0,
        spaceAfter=10
    )
    
    normal_style = ParagraphStyle(
        'TextoPersonalizado',
        parent=styles['Normal'],
        fontName=template.fonte_texto or 'Helvetica',
        fontSize=int(template.tamanho_texto or 12),
        alignment=0
    )
    
    # Elementos para o documento
    elementos = []
    
    # Adicionar logos do template
    if template.logo1_path or template.logo2_path or template.logo3_path:
        logos = []
        widths = []
        
        # Adicionar logos disponíveis
        for logo_path in [template.logo1_path, template.logo2_path, template.logo3_path]:
            if logo_path and os.path.exists(logo_path):
                img = Image(logo_path)
                img.drawHeight = 0.8 * inch
                img.drawWidth = 0.8 * inch
                logos.append(img)
                widths.append(2 * inch)
        
        if logos:
            tabela_logos = Table([logos], colWidths=widths)
            tabela_logos.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elementos.append(tabela_logos)
            elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar título
    titulo_texto = titulo or template.titulo_padrao or f"Relatório de Cadastros - {datetime.now().strftime('%d/%m/%Y')}"
    elementos.append(Paragraph(titulo_texto, titulo_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar cabeçalho do template
    if template.cabecalho:
        elementos.append(Paragraph(template.cabecalho, normal_style))
        elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar informações do relatório
    elementos.append(Paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}", normal_style))
    elementos.append(Paragraph(f"Total de registros: {len(cadastros)}", normal_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Tabela de dados (similar à função gerar_pdf_relatorio_formularios)
    cabecalho = []
    colunas_unicas = list(set(colunas))  # Garantir que as colunas não se repitam
    for coluna in colunas_unicas:
        if coluna == 'nome':
            cabecalho.append('Nome')
        elif coluna == 'projeto':
            cabecalho.append('Projeto')
        elif coluna == 'data_nascimento':
            cabecalho.append('Data Nasc.')
        elif coluna == 'data_cadastro':
            cabecalho.append('Data Cadastro')
        elif coluna == 'cras':
            cabecalho.append('CRAS')
        elif coluna == 'cidade':
            cabecalho.append('Cidade')
        elif coluna == 'bairro':
            cabecalho.append('Bairro')
        elif coluna == 'contato':
            cabecalho.append('Contato')
        elif coluna == 'nome_mae':
            cabecalho.append('Nome da Mãe')
    
    # Preparar dados da tabela
    dados_tabela = [cabecalho]
    for cadastro in cadastros:
        linha = []
        for coluna in colunas_unicas:  # Usar colunas únicas
            if coluna == 'nome':
                linha.append(cadastro.nome or '')
            elif coluna == 'projeto':
                linha.append(cadastro.projeto or '')
            elif coluna == 'data_nascimento':
                linha.append(cadastro.data_nascimento.strftime('%d/%m/%Y') if cadastro.data_nascimento else '')
            elif coluna == 'data_cadastro':
                linha.append(cadastro.data_cadastro.strftime('%d/%m/%Y') if cadastro.data_cadastro else '')
            elif coluna == 'cras':
                linha.append(cadastro.cras.nome if cadastro.cras else '')
            elif coluna == 'cidade':
                linha.append(cadastro.cidade or '')
            elif coluna == 'bairro':
                linha.append(cadastro.bairro or '')
            elif coluna == 'contato':
                linha.append(cadastro.contato or '')
            elif coluna == 'nome_mae':
                linha.append(cadastro.nome_mae or '')
        dados_tabela.append(linha)
    
    # Criar tabela
    tabela = Table(dados_tabela, repeatRows=1)
    
    # Estilo da tabela - adaptado de acordo com o template
    cor_cabecalho = template.cor_cabecalho or colors.lightgrey
    estilo_tabela = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), cor_cabecalho),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), template.fonte_texto or 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ])
    
    # Aplicar estilo alternado às linhas
    for i in range(1, len(dados_tabela)):
        if i % 2 == 0:
            estilo_tabela.add('BACKGROUND', (0, i), (-1, i), colors.lightgrey)
    
    tabela.setStyle(estilo_tabela)
    elementos.append(tabela)
    
    # Adicionar rodapé do template
    elementos.append(Spacer(1, 1 * cm))
    elementos.append(Paragraph("Relatório gerado pelo Sistema CRAS360 - © 2025", normal_style))
    
    # Gerar o documento
    doc.build(elementos)
    
    # Retornar o arquivo para download
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return django_send_file(arquivo_pdf, download_name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def gerar_doc_relatorio_formularios(cadastros, colunas, titulo=None):
    """
    Gera um relatório em formato DOCX a partir dos cadastros filtrados
    
    Args:
        cadastros: Lista de objetos Cadastro a serem incluídos no relatório
        colunas: Lista de colunas a serem exibidas
        titulo: Título opcional para o relatório
    
    Returns:
        Response com o arquivo DOCX para download
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Criar um arquivo temporário para o DOCX
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    arquivo_docx = temp_file.name
    temp_file.close()
    
    # Criar documento
    doc = Document()
    
    # Configurar margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # Adicionar título
    titulo_texto = titulo or f"Relatório de Cadastros - {datetime.now().strftime('%d/%m/%Y')}"
    heading = doc.add_heading(titulo_texto, level=1)
    
    # Adicionar informações do relatório
    p = doc.add_paragraph(f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p = doc.add_paragraph(f"Total de registros: {len(cadastros)}")
    doc.add_paragraph()  # Linha em branco
    
    # Criar tabela
    table = doc.add_table(rows=1, cols=len(colunas), style='Table Grid')
    
    # Configurar cabeçalho da tabela
    header_cells = table.rows[0].cells
    colunas_unicas = list(set(colunas))  # Garantir que as colunas não se repitam
    for i, coluna in enumerate(colunas_unicas):
        if coluna == 'nome':
            header_cells[i].text = 'Nome'
        elif coluna == 'projeto':
            header_cells[i].text = 'Projeto'
        elif coluna == 'data_nascimento':
            header_cells[i].text = 'Data Nasc.'
        elif coluna == 'data_cadastro':
            header_cells[i].text = 'Data Cadastro'
        elif coluna == 'cras':
            header_cells[i].text = 'CRAS'
        elif coluna == 'cidade':
            header_cells[i].text = 'Cidade'
        elif coluna == 'bairro':
            header_cells[i].text = 'Bairro'
        elif coluna == 'contato':
            header_cells[i].text = 'Contato'
        elif coluna == 'nome_mae':
            header_cells[i].text = 'Nome da Mãe'
    
    # Formatar cabeçalho
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Adicionar dados
    for cadastro in cadastros:
        row_cells = table.add_row().cells
        for i, coluna in enumerate(colunas_unicas):
            if coluna == 'nome':
                row_cells[i].text = cadastro.nome or ''
            elif coluna == 'projeto':
                row_cells[i].text = cadastro.projeto or ''
            elif coluna == 'data_nascimento':
                row_cells[i].text = cadastro.data_nascimento.strftime('%d/%m/%Y') if cadastro.data_nascimento else ''
            elif coluna == 'data_cadastro':
                row_cells[i].text = cadastro.data_cadastro.strftime('%d/%m/%Y') if cadastro.data_cadastro else ''
            elif coluna == 'cras':
                row_cells[i].text = cadastro.cras.nome if cadastro.cras else ''
            elif coluna == 'cidade':
                row_cells[i].text = cadastro.cidade or ''
            elif coluna == 'bairro':
                row_cells[i].text = cadastro.bairro or ''
            elif coluna == 'contato':
                row_cells[i].text = cadastro.contato or ''
            elif coluna == 'nome_mae':
                row_cells[i].text = cadastro.nome_mae or ''
    
    # Adicionar rodapé
    doc.add_paragraph()  # Linha em branco
    p = doc.add_paragraph("Relatório gerado pelo Sistema S.I.L.V.I.A. - © 2025")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    doc.save(arquivo_docx)
    
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return django_send_file(arquivo_docx, download_name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def gerar_doc_instrumental(titulo, conteudo=None):
    """
    Gera um documento DOCX instrumental com título e conteúdo
    
    Args:
        titulo: Título do documento
        conteudo: Conteúdo do documento (pode ser Delta JSON do Quill)
    
    Returns:
        Response com o arquivo DOCX para download
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Criar um arquivo temporário para o DOCX
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    arquivo_docx = temp_file.name
    temp_file.close()
    
    # Criar documento
    doc = Document()
    
    # Configurar margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
    
    # Adicionar título
    heading = doc.add_heading(titulo, level=1)
    
    # Processar conteúdo (pode ser Delta JSON do Quill)
    if conteudo:
        import json
        try:
            # Tentar interpretar como delta JSON
            delta = json.loads(conteudo)
            processar_delta_quill_docx(doc, delta)
        except:
            # Se falhar, tratar como texto simples
            paragrafos = conteudo.split('\n')
            for paragrafo in paragrafos:
                if paragrafo.strip():  # evitar adicionar parágrafos vazios
                    doc.add_paragraph(paragrafo)
    
    # Adicionar rodapé
    doc.add_paragraph()  # Linha em branco
    p = doc.add_paragraph(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    doc.save(arquivo_docx)
    
    # Retornar o arquivo para download
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return django_send_file(arquivo_docx, download_name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def gerar_doc_relatorio_com_template(cadastros, colunas, template_id, titulo=None):
    """
    Gera um relatório DOCX usando um template personalizado
    
    Args:
        cadastros: Lista de objetos Cadastro
        colunas: Lista de colunas a exibir
        template_id: ID do template a ser usado
        titulo: Título opcional para o relatório
    
    Returns:
        Response com o arquivo DOCX para download
    """
    # Implementação similar à função PDF, mas para DOCX
    # Para simplificar, vamos chamar a função padrão por enquanto
    return gerar_doc_relatorio_formularios(cadastros, colunas, titulo)

def gerar_excel_relatorio(cadastros, colunas):
    """
    Gera um relatório em formato Excel a partir dos cadastros filtrados
    
    Args:
        cadastros: Lista de objetos Cadastro a serem incluídos no relatório
        colunas: Lista de colunas a serem exibidas
    
    Returns:
        Response com o arquivo Excel para download
    """
    import pandas as pd
    
    # Criar um arquivo temporário para o Excel
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    arquivo_excel = temp_file.name
    temp_file.close()
    
    # Preparar dados para o DataFrame
    dados = []
    colunas_unicas = list(set(colunas))  # Garantir que as colunas não se repitam
    for cadastro in cadastros:
        linha = {}
        for coluna in colunas_unicas:
            if coluna == 'nome':
                linha['Nome'] = cadastro.nome or ''
            elif coluna == 'projeto':
                linha['Projeto'] = cadastro.projeto or ''
            elif coluna == 'data_nascimento':
                linha['Data Nascimento'] = cadastro.data_nascimento
            elif coluna == 'data_cadastro':
                linha['Data Cadastro'] = cadastro.data_cadastro
            elif coluna == 'cras':
                linha['CRAS'] = cadastro.cras.nome if cadastro.cras else ''
            elif coluna == 'cidade':
                linha['Cidade'] = cadastro.cidade or ''
            elif coluna == 'bairro':
                linha['Bairro'] = cadastro.bairro or ''
            elif coluna == 'contato':
                linha['Contato'] = cadastro.contato or ''
            elif coluna == 'nome_mae':
                linha['Nome da Mãe'] = cadastro.nome_mae or ''
        dados.append(linha)
    
    # Criar DataFrame e exportar para Excel
    df = pd.DataFrame(dados)
    df.to_excel(arquivo_excel, index=False)
    
    # Retornar o arquivo para download
    download_name = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return django_send_file(arquivo_excel, download_name, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# Funções auxiliares para processar Delta JSON do Quill
def processar_delta_quill(delta):
    """
    Processa um objeto delta do Quill e retorna uma lista de elementos ReportLab
    
    Args:
        delta: Objeto delta do Quill
    
    Returns:
        Lista de elementos ReportLab
    """
    elementos = []
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    
    # Processar cada operação
    for op in delta.get('ops', []):
        insert = op.get('insert', '')
        atributos = op.get('attributes', {})
        
        # Verificar se é uma imagem
        if isinstance(insert, dict) and 'image' in insert:
            try:
                # Processar a imagem
                img_data = insert['image']
                
                # Se a imagem for base64
                if img_data.startswith('data:image'):
                    # Extrair os dados da imagem
                    formato, dados_codificados = img_data.split(';base64,')
                    image_data = base64.b64decode(dados_codificados)
                    
                    # Criar imagem temporária
                    img_temp = BytesIO(image_data)
                    img = Image(img_temp, width=4*inch, height=3*inch)
                    elementos.append(img)
                    elementos.append(Spacer(1, 0.25 * cm))
            except Exception as e:
                print(f"Erro ao processar imagem: {str(e)}")
                continue
        
        # Processar texto
        elif isinstance(insert, str):
            # Quebrar texto em parágrafos
            paragrafos = insert.split('\n')
            for paragrafo in paragrafos:
                if paragrafo.strip() or len(paragrafos) > 1:  # incluir parágrafos vazios se houver quebra de linha
                    # Criar estilo com base nos atributos
                    estilo = ParagraphStyle(
                        'estilo_dinamico',
                        parent=normal_style,
                        fontName='Helvetica-Bold' if atributos.get('bold') else 'Helvetica',
                        fontSize=12,
                        leading=14
                    )
                    
                    # Aplicar itálico
                    if atributos.get('italic'):
                        if estilo.fontName == 'Helvetica-Bold':
                            estilo.fontName = 'Helvetica-BoldOblique'
                        else:
                            estilo.fontName = 'Helvetica-Oblique'
                    
                    # Aplicar alinhamento
                    if atributos.get('align') == 'center':
                        estilo.alignment = 1  # Centralizado
                    elif atributos.get('align') == 'right':
                        estilo.alignment = 2  # Direita
                    elif atributos.get('align') == 'justify':
                        estilo.alignment = 4  # Justificado
                    
                    elementos.append(Paragraph(paragrafo, estilo))
                    elementos.append(Spacer(1, 0.25 * cm))
    
    return elementos

def processar_delta_quill_docx(doc, delta):
    """
    Processa um objeto delta do Quill e adiciona o conteúdo ao documento DOCX
    
    Args:
        doc: Documento DOCX
        delta: Objeto delta do Quill
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import tempfile
    import os
    
    # Processar cada operação
    for op in delta.get('ops', []):
        insert = op.get('insert', '')
        atributos = op.get('attributes', {})
        
        # Verificar se é uma imagem
        if isinstance(insert, dict) and 'image' in insert:
            try:
                # Processar a imagem
                img_data = insert['image']
                
                # Se a imagem for base64
                if img_data.startswith('data:image'):
                    # Extrair os dados da imagem
                    formato, dados_codificados = img_data.split(';base64,')
                    image_data = base64.b64decode(dados_codificados)
                    
                    # Criar arquivo temporário para a imagem
                    img_temp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    img_temp.write(image_data)
                    img_temp.close()
                    
                    # Adicionar imagem ao documento
                    p = doc.add_paragraph()
                    r = p.add_run()
                    r.add_picture(img_temp.name, width=Inches(5))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Remover arquivo temporário
                    os.unlink(img_temp.name)
            except Exception as e:
                print(f"Erro ao processar imagem: {str(e)}")
                continue
        
        # Processar texto
        elif isinstance(insert, str):
            # Quebrar texto em parágrafos
            paragrafos = insert.split('\n')
            for paragrafo in paragrafos:
                if paragrafo.strip() or len(paragrafos) > 1:  # incluir parágrafos vazios se houver quebra de linha
                    p = doc.add_paragraph()
                    run = p.add_run(paragrafo)
                    
                    # Aplicar formatação
                    run.bold = atributos.get('bold', False)
                    run.italic = atributos.get('italic', False)
                    run.underline = atributos.get('underline', False)
                    
                    # Aplicar alinhamento
                    if atributos.get('align') == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif atributos.get('align') == 'right':
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif atributos.get('align') == 'justify':
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def gerar_pdf_ficha_paif(ficha):
    """
    Função especializada para gerar PDF de uma ficha PAIF
    com formatação específica
    """
    # Criar um arquivo temporário para o PDF
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    arquivo_pdf = temp_file.name
    temp_file.close()
    
    # Configurar o documento
    doc = SimpleDocTemplate(
        arquivo_pdf,
        pagesize=A4,
        leftMargin=MARGIN_LEFT,
        rightMargin=MARGIN_RIGHT,
        topMargin=MARGIN_TOP,
        bottomMargin=MARGIN_BOTTOM,
        title=f"Ficha PAIF - {ficha.numero_paif}"
    )
    
    # Preparar estilos
    styles = getSampleStyleSheet()
    titulo_style = styles['Heading1']
    subtitulo_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Elementos para o documento
    elementos = []
    
    # Cabeçalho
    elementos.append(Paragraph(f"FICHA PAIF Nº {ficha.numero_paif}", titulo_style))
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Dados principais
    elementos.append(Paragraph("DADOS DA FAMÍLIA", subtitulo_style))
    elementos.append(Spacer(1, 0.3 * cm))
    
    # Tabela de dados principais
    dados_principais = [
        ["Nome de Referência:", ficha.nome_referencia],
        ["CPF:", ficha.cpf],
        ["Data de Cadastro:", ficha.data.strftime("%d/%m/%Y") if hasattr(ficha, 'data') else ""],
        ["Endereço:", ficha.endereco],
        ["Telefone:", ficha.telefone]
    ]
    
    tabela_dados = Table(dados_principais, colWidths=[4*cm, 12*cm])
    tabela_dados.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ]))
    elementos.append(tabela_dados)
    elementos.append(Spacer(1, 0.5 * cm))
    
    # Composição familiar (se existir)
    if hasattr(ficha, 'familiares') and ficha.familiares.exists():
        elementos.append(Paragraph("COMPOSIÇÃO FAMILIAR", subtitulo_style))
        elementos.append(Spacer(1, 0.3 * cm))
        
        # Cabeçalho da tabela
        tabela_familiares_dados = [["Nome", "Parentesco", "Idade"]]
        
        # Linhas da tabela
        for familiar in ficha.familiares.all():
            tabela_familiares_dados.append([
                familiar.nome, 
                familiar.parentesco, 
                str(familiar.idade)
            ])
        
        tabela_familiares = Table(tabela_familiares_dados, colWidths=[8*cm, 5*cm, 3*cm])
        tabela_familiares.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ]))
        elementos.append(tabela_familiares)
        elementos.append(Spacer(1, 0.5 * cm))
    
    # Adicionar rodapé
    elementos.append(Spacer(1, 1 * cm))
    elementos.append(Paragraph(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}", normal_style))
    elementos.append(Paragraph("Sistema CRAS360 - © 2025", normal_style))
    
    # Gerar o documento
    doc.build(elementos)
    
    # Retornar o arquivo para download usando Django
    download_name = f"PAIF_{ficha.numero_paif}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return django_send_file(arquivo_pdf, download_name, 'application/pdf')

# Importe os módulos do Django
from django.http import HttpResponse
import os

# Função auxiliar para enviar arquivo para download no Django
def django_send_file(arquivo_path, download_name, mimetype):
    """
    Função auxiliar para enviar arquivo para download usando Django e excluir depois
    """
    with open(arquivo_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type=mimetype)
        response['Content-Disposition'] = f'attachment; filename="{download_name}"'
    
    # Excluir o arquivo temporário após a leitura
    os.unlink(arquivo_path)
    
    return response
